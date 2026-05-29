import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table):
    """Applies clean, professional borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="888888"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def parse_runs(paragraph_obj, text, base_font_size=11, is_italic=False, is_bold=False):
    """Parses markdown bold (**) and italic (*) syntax and appends appropriate runs."""
    # Pattern to capture bold, italic, or plain text chunks
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|[^*]+)')
    matches = pattern.findall(text)
    
    for match in matches:
        run = paragraph_obj.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(base_font_size)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # Determine formatting
        if match.startswith('***') and match.endswith('***'):
            run.text = match[3:-3]
            run.bold = True
            run.italic = True
        elif match.startswith('**') and match.endswith('**'):
            run.text = match[2:-2]
            run.bold = True
            run.italic = is_italic
        elif match.startswith('*') and match.endswith('*'):
            run.text = match[1:-1]
            run.italic = True
            run.bold = is_bold
        else:
            run.text = match
            run.bold = is_bold
            run.italic = is_italic

def compile_markdown_to_docx(md_path, docx_path, is_cover_letter=False):
    """Compiles a markdown file to a highly formatted native Word document."""
    print(f"Compiling {md_path} -> {docx_path}...")
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist!")
        return

    doc = Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    in_quote = False
    quote_text = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip('\n')
        
        # Handle Table parsing
        if line.strip().startswith('|'):
            in_table = True
            row_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # Skip separator line (e.g. |---|---|)
            if all(re.match(r'^:?-+:?$', c) for c in row_cells) and row_cells:
                i += 1
                continue
                
            if not table_headers:
                table_headers = row_cells
            else:
                table_rows.append(row_cells)
            i += 1
            continue
        elif in_table:
            # Table ended, write table to word doc
            if table_headers:
                cols = len(table_headers)
                table = doc.add_table(rows=1, cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                # Format Header
                hdr_cells = table.rows[0].cells
                for col_idx, text in enumerate(table_headers):
                    hdr_cells[col_idx].text = ""
                    p = hdr_cells[col_idx].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    parse_runs(p, text, base_font_size=10, is_bold=True)
                    set_cell_background(hdr_cells[col_idx], "F2F2F2")
                    
                # Format Data Rows
                for row_data in table_rows:
                    row_cells = table.add_row().cells
                    # Handle rows that might have fewer/more cells than header due to formatting
                    for col_idx in range(min(cols, len(row_data))):
                        row_cells[col_idx].text = ""
                        p = row_cells[col_idx].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after = Pt(3)
                        parse_runs(p, row_data[col_idx], base_font_size=9.5)
                        
                doc.add_paragraph().paragraph_format.space_before = Pt(6)
                
            in_table = False
            table_headers = []
            table_rows = []
            # Do not increment i, parse current line in next loop
            continue
            
        # Handle Alerts/Quotes (e.g., > [!IMPORTANT] or > )
        if line.strip().startswith('>'):
            in_quote = True
            # Strip the > and clean the text
            clean_line = line.strip().lstrip('>').strip()
            # Remove GitHub Alert headers like [!IMPORTANT], [!NOTE], etc.
            clean_line = re.sub(r'^\[!(IMPORTANT|NOTE|WARNING|CAUTION|TIP)\]', '', clean_line).strip()
            if clean_line:
                quote_text.append(clean_line)
            i += 1
            continue
        elif in_quote:
            # Quote ended, write quote block
            if quote_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                
                # Custom light grey background callout
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9F9F9"/>')
                pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="0086B3"/></w:pBdr>')
                pPr.append(shd)
                pPr.append(pbdr)
                
                full_quote = " ".join(quote_text)
                parse_runs(p, full_quote, base_font_size=10, is_italic=True)
                
            in_quote = False
            quote_text = []
            continue

        # Handle Headings and paragraphs
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
            
        if stripped.startswith('# '):
            # Document Title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(18)
            parse_runs(p, stripped[2:], base_font_size=18, is_bold=True)
        elif stripped.startswith('## '):
            # Heading 1
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            parse_runs(p, stripped[3:], base_font_size=14, is_bold=True)
        elif stripped.startswith('### '):
            # Heading 2
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            parse_runs(p, stripped[4:], base_font_size=12, is_bold=True)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            # List Item
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            parse_runs(p, stripped[2:])
        else:
            # Regular Paragraph
            p = doc.add_paragraph()
            if is_cover_letter:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(8)
            else:
                # Nature standard double or 1.5 line spacing
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(10)
                p.paragraph_format.space_before = Pt(0)
            
            # Align right for header meta block in cover letters
            if is_cover_letter and (i < 5 or stripped.startswith("Sincerely") or "dudekula" in stripped.lower()):
                if not stripped.startswith("Dear") and not stripped.startswith("Subject"):
                    if i < 4 or stripped.startswith("Sincerely") or "outlook.com" in stripped:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i < 4 else WD_ALIGN_PARAGRAPH.LEFT
            
            parse_runs(p, stripped)
            
        i += 1
        
    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")

if __name__ == "__main__":
    compile_markdown_to_docx("cover_letter.md", "cover_letter.docx", is_cover_letter=True)
    compile_markdown_to_docx("manuscript.md", "manuscript.docx")
    compile_markdown_to_docx("supplementary_material.md", "supplementary_material.docx")
    print("\nAll submission Word documents (.docx) compiled successfully!")
